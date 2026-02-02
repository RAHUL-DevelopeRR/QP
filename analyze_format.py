from docx import Document
from docx.shared import Inches, Pt, Cm

doc = Document(r'c:\Users\rahul\Downloads\QP\academicgen-engine\CIA_ACTUAL FORMAT.docx')

print("=== MARGINS ===")
section = doc.sections[0]
print(f"Top: {section.top_margin.inches} inches")
print(f"Bottom: {section.bottom_margin.inches} inches")
print(f"Left: {section.left_margin.inches} inches")
print(f"Right: {section.right_margin.inches} inches")

print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i}: [{p.alignment}] '{p.text}'")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    print(f"  Alignment: {table.alignment}")
    
    # Check for images in cells
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            # Check for images
            for para in cell.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//a:blip'):
                        print(f"  Row {ri}, Col {ci}: HAS IMAGE")
            
            text = cell.text.strip()[:60] if cell.text.strip() else "(empty)"
            if ri < 5:  # Only first 5 rows
                print(f"  Row {ri}, Col {ci}: {text}")
